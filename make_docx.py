"""Generate capstone.docx — run once, then commit the output file."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h


def add_para(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p


doc = Document()

# ── Title ────────────────────────────────────────────────────────────────────
title = doc.add_heading("COVID-19 Global Dashboard — Capstone Project", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# ── Section 1: Overview ───────────────────────────────────────────────────────
add_heading(doc, "1. Project Overview", level=1)
doc.add_paragraph(
    "This capstone project is an interactive web application built with Streamlit and "
    "Plotly that visualises COVID-19 time-series data sourced directly from the Johns "
    "Hopkins University Center for Systems Science and Engineering (JHU CSSE) public "
    "GitHub repository. Users can explore confirmed case counts and deaths for any "
    "combination of countries, toggle between daily new counts and cumulative totals, "
    "apply a 7-day rolling average to smooth noise, restrict the date range, and switch "
    "between line, bar, and area chart styles — all without writing a single line of code."
)

# ── Section 2: Data Source ────────────────────────────────────────────────────
add_heading(doc, "2. Data Source", level=1)
doc.add_paragraph(
    "Data is fetched at runtime (cached for 1 hour) from the JHU CSSE COVID-19 "
    "time-series files:"
)
for url in [
    "time_series_covid19_confirmed_global.csv",
    "time_series_covid19_deaths_global.csv",
]:
    doc.add_paragraph(
        f"https://raw.githubusercontent.com/CSSEGISandData/COVID-19/master/"
        f"csse_covid_19_data/csse_covid_19_time_series/{url}",
        style="List Bullet",
    )
doc.add_paragraph(
    "Province/State rows are aggregated to the country level by summing. "
    "Daily counts are derived by differencing consecutive cumulative values; "
    "negative values (data corrections) are clipped to zero."
)

# ── Section 3: Features ───────────────────────────────────────────────────────
add_heading(doc, "3. Features", level=1)
features = [
    "Multi-country selection — add or remove any country via a sidebar multiselect widget.",
    "Metric toggle — switch between Confirmed Cases and Deaths.",
    "Count type — choose Daily New counts or Cumulative totals.",
    "7-day rolling average — optional smoothing for daily series.",
    "Date range filter — restrict the view to any sub-period of the dataset.",
    "Chart type — Line, Bar, or Area charts powered by Plotly.",
    "Summary table — total confirmed, total deaths, and case-fatality rate per selected country.",
]
for f in features:
    doc.add_paragraph(f, style="List Bullet")

# ── Section 4: How to Run ─────────────────────────────────────────────────────
add_heading(doc, "4. How to Run Locally", level=1)
doc.add_paragraph("Prerequisites: Python 3.10+")
steps = [
    "Clone the repository:  git clone https://github.com/utkarshsharma19/covid-dashboard.git",
    "Enter the directory:   cd covid-dashboard",
    "Install dependencies:  pip install -r requirements.txt",
    "Launch the app:        streamlit run capstone.py",
    "Open your browser at   http://localhost:8501",
]
for s in steps:
    doc.add_paragraph(s, style="List Number")

# ── Section 5: Links ──────────────────────────────────────────────────────────
add_heading(doc, "5. Links", level=1)
links = [
    ("GitHub Repository", "https://github.com/utkarshsharma19/covid-dashboard"),
    ("Live Application (Streamlit Community Cloud)",
     "https://utkarshsharma19-covid-dashboard.streamlit.app"),
    ("Presentation (Google Slides / YouTube)",
     "https://github.com/utkarshsharma19/covid-dashboard#presentation"),
]
for label, url in links:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(f"{label}: ").bold = True
    p.add_run(url)

# ── Section 6: Tech Stack ─────────────────────────────────────────────────────
add_heading(doc, "6. Technology Stack", level=1)
stack = [
    "Python 3.10",
    "Streamlit 1.32+ — reactive web-app framework",
    "Plotly 5+ — interactive charting",
    "Pandas 2+ — data wrangling",
    "Requests — HTTP data fetching",
    "JHU CSSE COVID-19 Dataset — primary data source",
]
for s in stack:
    doc.add_paragraph(s, style="List Bullet")

# ── Footer ────────────────────────────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Author: Harsimrat  |  Course Capstone Project  |  2024")
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
run.font.size = Pt(9)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save("capstone.docx")
print("capstone.docx written successfully.")
